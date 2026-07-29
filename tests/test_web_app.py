"""Web 控制台端到端（离线）：TestClient 跑 scan-local job → 轮询至成功 → 台账新增。"""
import json
import time
from urllib.parse import parse_qs, urlparse

import pytest
from docx import Document
from fastapi.testclient import TestClient


@pytest.fixture
def client(tmp_path, monkeypatch):
    # 隔离到临时目录，避免动到真实 data/
    monkeypatch.setenv("KBM_LEDGER_DB", str(tmp_path / "ledger.db"))
    monkeypatch.setenv("KBM_WORK_DIR", str(tmp_path / "work"))
    monkeypatch.setenv("KBM_JOBS_DB", str(tmp_path / "web_jobs.db"))
    monkeypatch.setenv("KBM_TAXONOMY_FILE", "config/taxonomy.yaml")
    monkeypatch.setenv("KBM_FEISHU_TARGETS_FILE", str(tmp_path / "targets.json"))
    # Web 测试必须保持离线，避免项目 .env 中的真实凭证触发外部请求。
    monkeypatch.setenv("ANTHROPIC_API_KEY", "")
    monkeypatch.setenv("MS_TENANT_ID", "")
    monkeypatch.setenv("MS_CLIENT_ID", "")
    monkeypatch.setenv("MS_CLIENT_SECRET", "")
    from kb_migrator.config import get_settings
    get_settings.cache_clear()
    from kb_migrator.web.app import app
    yield TestClient(app)
    get_settings.cache_clear()


def _poll(client, job_id, timeout=15.0):
    deadline = time.time() + timeout
    while time.time() < deadline:
        j = client.get(f"/api/jobs/{job_id}").json()
        if j["status"] in ("success", "error"):
            return j
        time.sleep(0.1)
    raise AssertionError("job 超时未结束")


def test_status_and_settings_endpoints(client):
    s = client.get("/api/status").json()
    assert s["total"] == 0                     # 新台账为空
    assert "sink_ratio" in s
    assert "wecom_chat_ready" in s
    assert s["claude_ready"] is False
    assert s["claude_health"]["status"] == "missing_key"
    assert s["sharepoint_ready"] is False
    assert s["sharepoint_health"]["status"] == "missing_config"
    assert {
        "corp_id", "archive_secret", "private_key_file", "sdk_library",
    } == set(s["wecom_chat_checks"])
    fields = client.get("/api/settings").json()["fields"]
    assert "FEISHU_APP_ID" in fields and "ANTHROPIC_API_KEY" in fields


def test_ai_health_force_check_and_pipeline_preflight(client, monkeypatch):
    monkeypatch.setenv("ANTHROPIC_API_KEY", "sk-test-health")
    from kb_migrator.config import get_settings
    from kb_migrator.web import app as web_app

    get_settings.cache_clear()
    web_app._AI_HEALTH_CACHE.clear()
    quota = {
        "ready": False,
        "status": "quota",
        "label": "额度不足",
        "message": "AI 接口额度已用尽或预算上限已达到，请充值或提高预算后重试。",
    }
    monkeypatch.setattr(web_app, "_probe_ai_health", lambda settings: dict(quota))

    checked = client.post("/api/ai/health", json={"force": True})
    assert checked.status_code == 200
    assert checked.json()["status"] == "quota"

    status = client.get("/api/status").json()
    assert status["claude_ready"] is False
    assert status["claude_health"]["label"] == "额度不足"

    blocked = client.post("/api/jobs/pipeline", json={})
    assert blocked.status_code == 503
    assert blocked.json()["ai_health"]["status"] == "quota"
    assert "额度" in blocked.json()["error"]


def test_sharepoint_health_is_reflected_in_overview_status(client, monkeypatch):
    monkeypatch.setenv("MS_TENANT_ID", "tenant")
    monkeypatch.setenv("MS_CLIENT_ID", "client")
    monkeypatch.setenv("MS_CLIENT_SECRET", "secret")
    from kb_migrator.config import get_settings
    from kb_migrator.web import app as web_app

    get_settings.cache_clear()
    web_app._SHAREPOINT_HEALTH_CACHE.clear()
    permission_error = {
        "ready": False,
        "status": "permission",
        "label": "权限不足",
        "message": "缺少 Sites.Read.All/Files.Read.All 管理员授权。",
    }
    monkeypatch.setattr(
        web_app,
        "_probe_sharepoint_health",
        lambda settings: dict(permission_error),
    )

    checked = client.post("/api/sharepoint/health", json={"force": True})
    assert checked.status_code == 200
    assert checked.json()["status"] == "permission"

    status = client.get("/api/status").json()
    assert status["sharepoint_ready"] is False
    assert status["sharepoint_health"]["label"] == "权限不足"


def test_status_counts_wiki_target_bindings(client, tmp_path, monkeypatch):
    targets = tmp_path / "targets.json"
    targets.write_text(
        '{"mode":"wiki","folder_map":{},'
        '"wiki_node_map":{"分类一":"node-1","分类一旧称":"node-1",'
        '"分类二":"node-2"}}',
        encoding="utf-8",
    )
    monkeypatch.setenv("KBM_FEISHU_TARGETS_FILE", str(targets))
    from kb_migrator.config import get_settings
    get_settings.cache_clear()

    status = client.get("/api/status").json()

    assert status["targets_mode"] == "wiki"
    assert status["targets_count"] == 2


def test_active_wiki_structure_uses_drive_staging_for_load(client):
    from kb_migrator.config import get_settings
    from kb_migrator.ledger import Ledger
    from kb_migrator.structure import StructureService
    from kb_migrator.taxonomy import Taxonomy
    from kb_migrator.web.app import _load_routing_context

    settings = get_settings()
    tx = Taxonomy.load(settings.taxonomy_file)
    wiki_map = {
        path: f"wiki-{index}"
        for index, path in enumerate(tx.all_folder_paths(), 1)
    }
    targets = {
        "mode": "wiki",
        "root_token": "drive-staging-root",
        "space_id": "wiki-space",
        "folder_map": {},
        "wiki_node_map": wiki_map,
    }
    with open(settings.feishu_targets_file, "w", encoding="utf-8") as f:
        json.dump(targets, f, ensure_ascii=False)

    led = Ledger(settings.ledger_db)
    try:
        structures = StructureService(
            led, tx, settings.feishu_targets_file
        )
        active = structures.active_version()
        folders, node_ids, resolver, mode = _load_routing_context(
            structures, targets, active["id"]
        )
        category = tx.all_folder_paths()[0]
        resolved = resolver({"category": category})

        assert mode == "wiki"
        assert folders[category] == "drive-staging-root"
        assert node_ids[category]
        assert resolved["remote_token"] == "drive-staging-root"
        assert resolved["node_id"] == node_ids[category]
    finally:
        led.close()


def test_runtime_handshake_readiness_and_no_cache_headers(client):
    meta_response = client.get("/api/meta")
    assert meta_response.status_code == 200
    meta = meta_response.json()
    assert meta["api_protocol"] == 1
    assert meta["app_version"]
    assert meta["instance_id"]
    assert {
        "structure_workbench", "structure_reconciliation", "item_relocation",
    }.issubset(meta["capabilities"])

    ready_response = client.get("/api/health/ready")
    assert ready_response.status_code == 200
    ready = ready_response.json()
    assert ready["ready"] is True
    assert ready["checks"]["ledger"]["quick_check"] == "ok"
    assert ready["checks"]["ledger"]["missing_tables"] == []
    assert ready["checks"]["taxonomy"]["categories"] > 0

    for response in (client.get("/"), meta_response, ready_response):
        assert response.headers["cache-control"].startswith("no-store")
        assert response.headers["x-kbm-api-protocol"] == "1"
        assert response.headers["x-kbm-instance-id"] == meta["instance_id"]


def test_frontend_navigation_and_layout_controls(client):
    html = client.get("/").text
    assert html.index('data-tab="auth"') < html.index('data-tab="targets"')
    assert html.index('data-tab="targets"') < html.index('data-tab="migrate"')
    assert "全部展开" in html and "全部收起" in html
    assert 'class="scroll-window"' in html
    assert "migrationFlowButton('sources'" in html
    assert "migrationFlowButton('process'" in html
    assert "migrationFlowButton('delivery'" in html
    assert "[hidden]{display:none!important}" in html
    assert 'class="migration-flow"' in html
    assert "migrationFlowState" in html
    assert "tab.disabled=!state.enabled" in html
    assert 'class="governance-metrics"' in html
    assert "governance-scroll" in html
    assert "企业微信群聊（会话存档）" in html
    assert "${dot(s.sharepoint_ready)}SharePoint" in html
    assert "refreshSharePointHealth" in html
    assert html.count("测试企业微信凭证") == 1
    assert "kbm:migration-inputs:v1" in html
    assert "kbm:last-migration-job:v1" in html
    assert "restoreMigrationInputs" in html
    assert "restoreMigrationJob" in html
    assert "清空输入" in html and "清空进度" in html
    assert "align-items:stretch" in html
    assert ".migration-monitor{position:static;margin:0;height:100%" in html


def test_frontend_read_endpoints_return_expected_shapes(client):
    assert "items" in client.get("/api/review").json()
    assert "items" in client.get("/api/failures").json()
    assert "calibration" in client.get("/api/insights").json()
    assert "review_due" in client.get("/api/governance").json()
    assert "summary" in client.get("/api/targets").json()
    draft = client.post("/api/structures/drafts", json={}).json()["structure"]
    assert draft["status"] == "draft"
    assert client.get(f"/api/structures/drafts/{draft['id']}").status_code == 200
    assert "jobs" in client.get("/api/jobs").json()


def test_single_and_batch_failure_retry_only_requeues_selected(client):
    from kb_migrator.config import get_settings
    from kb_migrator.ledger import Ledger
    from kb_migrator.models import SourceItem, SourceType, Stage

    led = Ledger(get_settings().ledger_db)
    keys = []
    try:
        for index, failed_stage in enumerate(("classify", "classify", "fetch"), 1):
            item = SourceItem(
                source_type=SourceType.LOCAL,
                source_id=f"retry-{index}",
                source_path=f"/tmp/retry-{index}.docx",
                original_name=f"重试-{index}.docx",
            )
            led.upsert_discovered(item)
            key = item.stable_key()
            led.mark_failed(key, failed_stage, "测试失败")
            keys.append(key)
    finally:
        led.close()

    single = client.post("/api/failures/retry", json={"keys": [keys[0]]})
    assert single.status_code == 200
    assert single.json()["requeued_count"] == 1
    assert "去重与分类" in single.json()["next_actions"][0]

    batch = client.post("/api/failures/retry", json={"keys": keys[1:]})
    assert batch.status_code == 200
    assert batch.json()["requeued_count"] == 2
    assert len(batch.json()["next_actions"]) == 2

    led = Ledger(get_settings().ledger_db)
    try:
        assert led.get(keys[0])["stage"] == Stage.DEDUPED.value
        assert led.get(keys[1])["stage"] == Stage.DEDUPED.value
        assert led.get(keys[2])["stage"] == Stage.DISCOVERED.value
        assert all(led.get(key)["error_detail"] is None for key in keys)
    finally:
        led.close()


def test_structure_draft_save_snapshot_and_diff_api(client):
    draft = client.post("/api/structures/drafts", json={}).json()["structure"]
    first = draft["nodes"][0]
    first_id = first["node_id"]
    first["display_name"] = "制度流程中心"
    saved = client.put(f"/api/structures/drafts/{draft['id']}", json={
        "revision": draft["revision"],
        "name": draft["name"],
        "root_name": draft["root_name"],
        "nodes": draft["nodes"],
    })
    assert saved.status_code == 200
    body = saved.json()["structure"]
    assert body["nodes"][0]["node_id"] == first_id
    assert "01 制度与流程" in body["nodes"][0]["aliases"]

    snapshot = client.post("/api/remote-structures/refresh", json={
        "mode": "drive",
        "root_token": "root",
        "nodes": [{
            "remote_token": "existing",
            "parent_token": "root",
            "display_name": "制度流程中心",
            "node_type": "folder",
        }],
    })
    assert snapshot.status_code == 200
    snapshot_id = snapshot.json()["snapshot"]["id"]
    diff = client.post(
        f"/api/structures/drafts/{draft['id']}/diff",
        json={"snapshot_id": snapshot_id},
    )
    assert diff.status_code == 200
    assert "MAP" in {
        action["action_type"] for action in diff.json()["plan"]["actions"]
    }


def test_remote_structure_mapping_adoption_decision_and_suggestions_api(client):
    draft = client.post("/api/structures/drafts", json={}).json()["structure"]
    snapshot = client.post("/api/remote-structures/refresh", json={
        "mode": "drive", "root_token": "root",
        "nodes": [
            {
                "remote_token": "map-me", "parent_token": "root",
                "display_name": "已有目录", "file_count": 4,
            },
            {
                "remote_token": "adopt-me", "parent_token": "root",
                "display_name": "飞书新增目录", "file_count": 2,
            },
            {
                "remote_token": "external-me", "parent_token": "root",
                "display_name": "外部维护目录",
            },
        ],
    }).json()["snapshot"]
    mapped = client.post(
        f"/api/structures/drafts/{draft['id']}/map-remote",
        json={
            "revision": draft["revision"],
            "node_id": draft["nodes"][0]["node_id"],
            "remote_token": "map-me",
        },
    )
    assert mapped.status_code == 200
    current = mapped.json()["structure"]
    adopted = client.post(
        f"/api/structures/drafts/{draft['id']}/adopt-remote",
        json={"revision": current["revision"], "remote_token": "adopt-me"},
    )
    assert adopted.status_code == 200
    decision = client.post("/api/remote-structures/decisions", json={
        "mode": "drive", "remote_token": "external-me",
        "decision": "external", "note": "业务系统自行维护",
    })
    assert decision.status_code == 200
    latest = client.get(
        "/api/remote-structures/snapshots/" + snapshot["id"]
    ).json()["snapshot"]
    assert {
        node["remote_token"]: (node.get("decision") or {}).get("decision")
        for node in latest["nodes"]
    }["external-me"] == "external"
    assert client.get("/api/structures").json()["versions"]
    suggestions = client.post(
        f"/api/structures/drafts/{draft['id']}/generate", json={}
    )
    assert suggestions.status_code == 200
    assert suggestions.json()["notice"]


def test_structure_change_plan_scope_approval_and_cancel_api(client):
    draft = client.post("/api/structures/drafts", json={}).json()["structure"]
    plan = client.post(
        f"/api/structures/drafts/{draft['id']}/diff",
        json={"history_scope": "unmigrated_only", "actor": "planner"},
    ).json()["plan"]
    assert plan["status"] == "preview"
    assert all("action_order" in action for action in plan["actions"])
    assert all("depends_on" in action for action in plan["actions"])

    updated = client.put(f"/api/structure-plans/{plan['id']}", json={
        "revision": plan["revision"],
        "history_scope": "include_retries",
        "actor": "planner",
    })
    assert updated.status_code == 200
    plan = updated.json()["plan"]
    assert plan["history_scope"] == "include_retries"

    approved_version = client.post(
        f"/api/structures/drafts/{draft['id']}/approve",
        json={"actor": "owner"},
    )
    assert approved_version.status_code == 200
    approved_plan = client.post(
        f"/api/structure-plans/{plan['id']}/approve",
        json={"actor": "owner"},
    )
    assert approved_plan.status_code == 200
    assert approved_plan.json()["plan"]["status"] == "approved"
    impact = client.get(f"/api/structure-plans/{plan['id']}/impact")
    assert impact.status_code == 200
    assert impact.json()["history_scope"] == "include_retries"
    cancelled = client.post(
        f"/api/structure-plans/{plan['id']}/cancel", json={"actor": "owner"}
    )
    assert cancelled.status_code == 200
    assert cancelled.json()["plan"]["status"] == "cancelled"


def test_structure_approval_rejects_plan_from_stale_remote_snapshot(client):
    draft = client.post("/api/structures/drafts", json={}).json()["structure"]
    first = client.post("/api/remote-structures/refresh", json={
        "mode": draft["mode"], "nodes": [],
    }).json()["snapshot"]
    client.post(
        f"/api/structures/drafts/{draft['id']}/diff",
        json={"snapshot_id": first["id"]},
    )
    second = client.post("/api/remote-structures/refresh", json={
        "mode": draft["mode"], "nodes": [],
    }).json()["snapshot"]
    stale = client.post(
        f"/api/structures/drafts/{draft['id']}/approve",
        json={"actor": "owner"},
    )
    assert stale.status_code == 409
    assert "最新飞书快照" in stale.json()["error"]

    fresh_plan = client.post(
        f"/api/structures/drafts/{draft['id']}/diff",
        json={"snapshot_id": second["id"]},
    ).json()["plan"]
    approved = client.post(
        f"/api/structures/drafts/{draft['id']}/approve",
        json={"actor": "owner"},
    )
    assert approved.status_code == 200
    plan_approved = client.post(
        f"/api/structure-plans/{fresh_plan['id']}/approve",
        json={"actor": "release-owner"},
    )
    assert plan_approved.status_code == 200


def test_structure_merge_api_preserves_target_and_records_intent(client):
    draft = client.post("/api/structures/drafts", json={}).json()["structure"]
    target, source = draft["nodes"][0], draft["nodes"][1]
    response = client.post(
        f"/api/structures/drafts/{draft['id']}/merge",
        json={
            "revision": draft["revision"],
            "target_node_id": target["node_id"],
            "source_node_ids": [source["node_id"]],
        },
    )
    assert response.status_code == 200
    merged = response.json()["structure"]
    assert any(node["node_id"] == target["node_id"] for node in merged["nodes"])
    assert not any(node["node_id"] == source["node_id"] for node in merged["nodes"])
    assert source["display_name"] in next(
        node for node in merged["nodes"] if node["node_id"] == target["node_id"]
    )["aliases"]
    assert merged["transformations"][0]["transformation_type"] == "merge"
    audit = client.get(f"/api/structures/{draft['id']}/audit")
    assert audit.status_code == 200
    assert any(
        event["event_type"] == "nodes_merged"
        for event in audit.json()["events"]
    )


def test_structure_split_health_and_multi_approval_api(client):
    draft = client.post("/api/structures/drafts", json={}).json()["structure"]
    source = draft["nodes"][0]
    response = client.post(
        f"/api/structures/drafts/{draft['id']}/split",
        json={
            "revision": draft["revision"],
            "source_node_id": source["node_id"],
            "children": [
                {
                    "display_name": "年度制度",
                    "assignment_rule": {
                        "field": "year", "operator": "equals",
                        "value": "2026", "priority": 10,
                    },
                },
                {
                    "display_name": "其他制度",
                    "assignment_rule": {"fallback": True, "priority": 999},
                },
            ],
        },
    )
    assert response.status_code == 200
    split = response.json()["structure"]
    assert any(node["assignment_rule"] for node in split["nodes"])

    health = client.get(f"/api/structures/{draft['id']}/health")
    assert health.status_code == 200
    assert health.json()["validation"]["valid"] is True
    diff = client.post(
        f"/api/structures/drafts/{draft['id']}/diff", json={}
    )
    assert diff.status_code == 200

    first = client.post(
        f"/api/structures/drafts/{draft['id']}/approve",
        json={"actor": "alice", "required_approvals": 2},
    )
    assert first.status_code == 200
    assert first.json()["structure"]["status"] == "reviewing"
    second = client.post(
        f"/api/structures/drafts/{draft['id']}/approve",
        json={"actor": "bob", "required_approvals": 2},
    )
    assert second.status_code == 200
    assert second.json()["structure"]["status"] == "approved"


def test_historical_relocation_plan_api(client):
    from kb_migrator.config import get_settings
    from kb_migrator.ledger import Ledger
    from kb_migrator.models import SourceItem, SourceType, Stage
    from kb_migrator.structure import StructureService
    from kb_migrator.taxonomy import Taxonomy

    settings = get_settings()
    led = Ledger(settings.ledger_db)
    structures = StructureService(
        led, Taxonomy.load(settings.taxonomy_file),
        settings.feishu_targets_file,
    )
    try:
        draft = structures.ensure_draft()
        source = draft["nodes"][0]
        split = structures.split_node(
            draft["id"], draft["revision"], source["node_id"], [
                {
                    "display_name": "PDF 文档",
                    "assignment_rule": {
                        "field": "doc_type", "operator": "equals",
                        "value": "pdf", "priority": 10,
                    },
                },
                {
                    "display_name": "其他文档",
                    "assignment_rule": {"fallback": True, "priority": 999},
                },
            ],
        )
        for index, node in enumerate(split["nodes"]):
            structures.bind_node(
                split["id"], node["node_id"], "drive", f"folder-{index}"
            )
        structures.approve(split["id"])
        active = structures.activate(split["id"], root_token="root")
        item = SourceItem(
            source_type=SourceType.LOCAL, source_id="history-api",
            source_path="/history.pdf", original_name="history.pdf",
        )
        led.upsert_discovered(item)
        led.update(
            item.stable_key(), category=source["display_name"],
            stage=Stage.LOADED.value, feishu_token="file-history",
        )
        led.assign_structure_target(
            item.stable_key(), active["id"], source["node_id"]
        )
    finally:
        led.close()

    created = client.post(
        f"/api/structures/{active['id']}/relocation-plans",
        json={"actor": "planner"},
    )
    assert created.status_code == 200
    plan = created.json()["plan"]
    assert plan["summary"]["total"] == 1

    latest = client.get(
        f"/api/structures/{active['id']}/relocation-plans/latest"
    )
    assert latest.json()["plan"]["id"] == plan["id"]
    selected = client.put(
        f"/api/relocation-plans/{plan['id']}",
        json={
            "revision": plan["revision"],
            "stable_keys": [plan["actions"][0]["stable_key"]],
        },
    )
    assert selected.status_code == 200
    approved = client.post(
        f"/api/relocation-plans/{plan['id']}/approve",
        json={"actor": "approver"},
    )
    assert approved.status_code == 200
    assert approved.json()["plan"]["status"] == "approved"


def test_frontend_controls_reference_existing_api_routes(client):
    html = client.get("/").text
    for path in (
        "/api/meta", "/api/health/ready", "/api/status", "/api/settings",
        "/api/review", "/api/governance",
        "/api/insights", "/api/failures", "/api/structures/drafts",
        "/api/remote-structures/latest", "/api/remote-structures/refresh",
        "/api/jobs/scan-local", "/api/jobs/sharepoint", "/api/jobs/wedrive",
        "/api/jobs/wecom-chat", "/api/jobs/pipeline", "/api/jobs/semantic",
        "/api/jobs/govern-chat", "/api/jobs/structure-apply", "/api/jobs/load",
        "/api/jobs/relocation-apply", "/api/relocation-plans/",
        "/api/jobs/retry", "/api/jobs/push-to-wiki", "/api/jobs/archive",
    ):
        assert path in html
    assert "<a href=\"/api/oauth/feishu/login\"><button" not in html
    assert "CLIENT_API_PROTOCOL=1" in html
    assert "python console.py restart" in html


def test_confirm_and_reject_validate_item_and_category(client):
    missing = client.post("/api/confirm", json={
        "key": "local:not-found", "category": "01 制度与流程",
    })
    assert missing.status_code == 404
    assert client.post(
        "/api/reject", json={"key": "local:not-found"},
    ).status_code == 404


def test_oauth_state_is_bound_to_http_only_browser_cookie(client, monkeypatch):
    monkeypatch.setenv("FEISHU_APP_ID", "cli_test")
    monkeypatch.setenv("FEISHU_APP_SECRET", "secret_test")
    from kb_migrator.config import get_settings
    get_settings.cache_clear()

    login = client.get("/api/oauth/feishu/login", follow_redirects=False)
    assert login.status_code == 302
    assert "HttpOnly" in login.headers["set-cookie"]
    state = parse_qs(urlparse(login.headers["location"]).query)["state"][0]

    # 即使 state 仍在服务端白名单中，没有发起授权时的浏览器 Cookie 也必须拒绝。
    client.cookies.delete("kbm_oauth_state")
    callback = client.get(f"/feishu/oauth/callback?state={state}&code=fake")
    assert callback.status_code == 400
    assert "会话不匹配" in callback.text


def test_scan_local_job_end_to_end(client, tmp_path):
    src = tmp_path / "src"
    src.mkdir()
    d = Document()
    d.add_paragraph("远程办公管理制度")
    d.add_paragraph("本管理办法规定流程与审批。")
    d.save(str(src / "制度.docx"))

    r = client.post("/api/jobs/scan-local", json={"path": str(src)})
    assert r.status_code == 200
    job = _poll(client, r.json()["job_id"])
    assert job["status"] == "success", job
    assert job["result"]["discovered"] == 1
    assert job["result"]["extracted"] == 1

    # 台账已新增
    assert client.get("/api/status").json()["total"] == 1


def test_scan_local_rejects_bad_path(client):
    r = client.post("/api/jobs/scan-local", json={"path": "Z:/nonexistent-xyz"})
    assert r.status_code == 400


def test_retry_job_dry_run(client):
    # 空台账下 retry(dry-run) 应成功：重排 0 条、无真实写入
    r = client.post("/api/jobs/retry", json={"dry_run": True})
    assert r.status_code == 200
    job = _poll(client, r.json()["job_id"])
    assert job["status"] == "success", job
    assert job["result"]["requeued"] == 0


def test_structure_apply_job_dry_run(client):
    draft = client.post("/api/structures/drafts", json={}).json()["structure"]
    diff = client.post(
        f"/api/structures/drafts/{draft['id']}/diff", json={}
    )
    assert diff.status_code == 200
    approved = client.post(
        f"/api/structures/drafts/{draft['id']}/approve", json={}
    )
    assert approved.status_code == 200
    response = client.post("/api/jobs/structure-apply", json={
        "version_id": draft["id"], "dry_run": True,
    })
    job = _poll(client, response.json()["job_id"])
    assert job["status"] == "success", job
    assert job["result"]["planned_nodes"] == len(draft["nodes"])


@pytest.mark.parametrize("path", [
    "/api/jobs/load", "/api/jobs/push-to-wiki", "/api/jobs/archive",
])
def test_external_write_jobs_default_to_safe_dry_run(client, path):
    r = client.post(path, json={})
    assert r.status_code == 200
    job = _poll(client, r.json()["job_id"])
    assert job["status"] == "success", job


def test_sharepoint_job_without_creds_errors(client):
    # 无 MS_* 凭证时 job 应优雅失败（error），不触网
    r = client.post("/api/jobs/sharepoint", json={"site": ""})
    assert r.status_code == 200
    job = _poll(client, r.json()["job_id"])
    assert job["status"] == "error"
    assert "SharePoint" in (job["error"] or "")


def test_wecom_chat_job_degrades_when_sdk_absent(client):
    # 无原生存档 SDK 时不报错，返回 {ready: False}（降级仅迁群文件）
    r = client.post("/api/jobs/wecom-chat", json={"chat_id": "chatX"})
    assert r.status_code == 200
    job = _poll(client, r.json()["job_id"])
    assert job["status"] == "success", job
    assert job["result"]["ready"] is False
